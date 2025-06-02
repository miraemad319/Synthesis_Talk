# backend/routes/upload.py - COMPLETE FIXED VERSION

from fastapi import APIRouter, File, UploadFile, Cookie, Response, Query, HTTPException, Request
from fastapi.responses import JSONResponse
import uuid
import hashlib
from typing import Optional
import time
import logging
import io

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter()

# Simple in-memory storage - shared with chat system
simple_document_store = {}

def extract_text_from_memory(file_content: bytes, filename: str, file_ext: str) -> str:
    """
    Extract text directly from memory without creating temp files
    """
    try:
        if file_ext in ['txt', 'md', 'csv']:
            # Direct text extraction from bytes
            text = file_content.decode('utf-8', errors='ignore')
            logger.info(f"[EXTRACT] Text file decoded: {len(text)} characters")
            return text
            
        elif file_ext == 'pdf':
            # Try simple PDF extraction from memory
            try:
                import PyPDF2
                pdf_file = io.BytesIO(file_content)
                reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    logger.info(f"[EXTRACT] PDF page {page_num + 1}: {len(page_text)} chars")
                
                if text.strip():
                    logger.info(f"[EXTRACT] PDF extraction successful: {len(text)} characters")
                    return text
                else:
                    logger.warning("[EXTRACT] PDF extraction returned empty text")
                    return f"PDF document: {filename}\n\nThis PDF was uploaded successfully but text extraction returned no content. The file may contain images or be password protected."
                    
            except Exception as e:
                logger.warning(f"[EXTRACT] PDF extraction failed: {e}")
                return f"PDF document: {filename}\n\nThis PDF was uploaded successfully. Text extraction encountered an issue: {str(e)}\n\nYou can still reference this document in conversations."
                
        elif file_ext in ['docx', 'doc']:
            # Try simple DOCX extraction from memory
            try:
                from docx import Document
                docx_file = io.BytesIO(file_content)
                doc = Document(docx_file)
                
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                
                if text.strip():
                    logger.info(f"[EXTRACT] DOCX extraction successful: {len(text)} characters")
                    return text
                else:
                    logger.warning("[EXTRACT] DOCX extraction returned empty text")
                    return f"Word document: {filename}\n\nThis document was uploaded successfully but appears to be empty or contains only formatting."
                    
            except Exception as e:
                logger.warning(f"[EXTRACT] DOCX extraction failed: {e}")
                return f"Word document: {filename}\n\nThis document was uploaded successfully. Text extraction encountered an issue: {str(e)}\n\nYou can still reference this document in conversations."
        
        else:
            return f"Document: {filename}\n\nUploaded successfully. File type: {file_ext.upper()}"
            
    except Exception as e:
        logger.error(f"[EXTRACT] General extraction error: {e}")
        return f"Document: {filename}\n\nUploaded successfully but text extraction failed: {str(e)}"

def simple_chunk_text(text: str, max_length: int = 1000) -> list:
    """
    Simple text chunking without external dependencies
    """
    if not text or len(text) <= max_length:
        return [text] if text else [""]
    
    # Split by double newlines first (paragraphs)
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # If adding this paragraph would exceed max_length
        if len(current_chunk) + len(para) + 2 > max_length:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = para
            else:
                # Single paragraph is too long, split by sentences
                sentences = para.split('. ')
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) + 2 > max_length:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = sentence
                    else:
                        current_chunk += sentence + ". "
        else:
            current_chunk += para + "\n\n"
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks if chunks else [text[:max_length]]

def get_current_context(request: Request, session_id: str) -> str:
    """Get current context from request or default"""
    # Try to get from cookie first
    current_context = request.cookies.get('current_context', 'default')
    
    # If not in cookie, try to get from context system
    try:
        # Import here to avoid circular imports
        from backend.routes.context import _current_context
        if _current_context:
            current_context = _current_context
    except ImportError:
        pass
    
    return current_context

@router.post("/upload/")
async def upload_file(
    request: Request,
    response: Response,
    file: UploadFile = File(...),
    session_id: Optional[str] = Cookie(default=None),
    format: str = Query(default="paragraph", enum=["paragraph", "bullets"]),
    context_id: Optional[str] = Query(default=None),
    allow_duplicates: bool = Query(default=False, description="Allow uploading the same file to different contexts")
):
    """
    FIXED: Context-aware upload processing with duplicate handling options
    """
    start_time = time.time()
    logger.info(f"[UPLOAD] Starting context-aware upload for: {file.filename}")
    
    try:
        # 1) Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # 2) Get file extension
        file_ext = file.filename.split('.')[-1].lower() if '.' in file.filename else 'unknown'
        allowed_extensions = {'pdf', 'txt', 'docx', 'doc', 'md', 'csv'}
        
        if file_ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{file_ext}")
        
        logger.info(f"[UPLOAD] File validation OK: {file.filename} (.{file_ext})")
        
        # 3) Read file content directly into memory
        file_content = await file.read()
        file_size = len(file_content)
        
        if file_size > 10 * 1024 * 1024:  # 10MB limit
            raise HTTPException(status_code=413, detail="File too large (max 10MB)")
        
        if file_size == 0:
            raise HTTPException(status_code=400, detail="File is empty")
        
        logger.info(f"[UPLOAD] File read OK: {file_size} bytes")
        
        # 4) Create/get session
        if session_id is None:
            session_id = str(uuid.uuid4())
            is_new_session = True
        else:
            is_new_session = False
        
        # 5) Get current context
        current_context = context_id or get_current_context(request, session_id)
        logger.info(f"[UPLOAD] Using context: {current_context}")
        
        # 6) FIXED: Context-aware duplicate checking
        file_hash = hashlib.md5(file_content).hexdigest()
        
        if session_id not in simple_document_store:
            simple_document_store[session_id] = []
        
        existing_docs = simple_document_store[session_id]
        
        # Check for duplicates within the current context
        duplicate_in_context = None
        duplicate_in_other_contexts = []
        
        for doc in existing_docs:
            doc_context = doc.get('context_id', 'default')
            if doc.get('file_hash') == file_hash:
                if doc_context == current_context:
                    duplicate_in_context = doc
                else:
                    duplicate_in_other_contexts.append(doc)
        
        # Handle duplicates based on context
        if duplicate_in_context and not allow_duplicates:
            return JSONResponse(
                status_code=409,
                content={
                    "error": "duplicate_in_context",
                    "message": f"File '{file.filename}' already exists in context '{current_context}'",
                    "options": {
                        "overwrite": "Replace the existing file in this context",
                        "rename": "Rename this file and upload as new",
                        "switch_context": "Upload to a different context",
                        "allow": f"Upload anyway (add ?allow_duplicates=true)"
                    },
                    "existing_file": {
                        "filename": duplicate_in_context.get('filename'),
                        "context": current_context,
                        "upload_time": duplicate_in_context.get('upload_time', 'unknown'),
                        "file_size": duplicate_in_context.get('file_size', 0)
                    },
                    "other_contexts": [
                        {
                            "context": doc.get('context_id', 'default'),
                            "filename": doc.get('filename'),
                            "upload_time": doc.get('upload_time', 'unknown')
                        } for doc in duplicate_in_other_contexts
                    ]
                }
            )
        
        # If file exists in other contexts, inform user but allow upload
        if duplicate_in_other_contexts:
            logger.info(f"[UPLOAD] File exists in other contexts: {[doc.get('context_id') for doc in duplicate_in_other_contexts]}")
        
        logger.info(f"[UPLOAD] Session setup OK: {session_id}")
        
        # 7) Extract text directly from memory
        logger.info("[UPLOAD] Starting text extraction from memory...")
        text = extract_text_from_memory(file_content, file.filename, file_ext)
        logger.info(f"[UPLOAD] Text extraction complete: {len(text)} characters")
        
        # 8) Simple chunking
        logger.info("[UPLOAD] Starting text chunking...")
        chunks = simple_chunk_text(text, max_length=800)
        logger.info(f"[UPLOAD] Chunking complete: {len(chunks)} chunks")
        
        # 9) FIXED: Store with context information
        doc_info = {
            'filename': file.filename,
            'file_hash': file_hash,
            'context_id': current_context,  # Context-specific storage
            'chunks': chunks,
            'full_text': text,
            'file_size': file_size,
            'file_type': file_ext,
            'upload_time': time.time(),
            'chunk_count': len(chunks),
            'upload_timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'processing_time': time.time() - start_time
        }
        
        # Remove existing file in context if overwriting
        if duplicate_in_context and allow_duplicates:
            simple_document_store[session_id] = [
                doc for doc in simple_document_store[session_id] 
                if not (doc.get('file_hash') == file_hash and doc.get('context_id') == current_context)
            ]
            logger.info(f"[UPLOAD] Replaced existing file in context {current_context}")
        
        simple_document_store[session_id].append(doc_info)
        logger.info(f"[UPLOAD] Storage complete in context: {current_context}")
        
        # 10) Create response
        processing_time = time.time() - start_time
        
        result = {
            "filename": file.filename,
            "summary": f"Successfully processed {file.filename} in context '{current_context}'. Extracted text and created {len(chunks)} chunks for analysis.",
            "chunks": len(chunks),
            "file_size": file_size,
            "session_id": session_id,
            "context_id": current_context,
            "file_type": file_ext,
            "success": True,
            "processing_time": f"{processing_time:.2f}s",
            "text_preview": text[:200] + "..." if len(text) > 200 else text,
            "extraction_method": "context_aware_memory",
            "duplicate_info": {
                "exists_in_other_contexts": len(duplicate_in_other_contexts) > 0,
                "other_contexts": [doc.get('context_id') for doc in duplicate_in_other_contexts],
                "replaced_existing": duplicate_in_context is not None and allow_duplicates
            }
        }
        
        logger.info(f"[UPLOAD] SUCCESS: {file.filename} processed in {processing_time:.2f}s for context {current_context}")
        
        json_response = JSONResponse(content=result)
        if is_new_session:
            json_response.set_cookie(key="session_id", value=session_id, httponly=True)
        
        # Update current context cookie
        json_response.set_cookie(key="current_context", value=current_context, httponly=True)
        
        return json_response
        
    except HTTPException:
        raise
    except Exception as e:
        processing_time = time.time() - start_time
        logger.error(f"[UPLOAD] FAILED after {processing_time:.2f}s: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.get("/upload/test")
async def test_upload_endpoint():
    """Test endpoint"""
    return {
        "status": "healthy",
        "message": "Context-aware upload endpoint",
        "temp_files_created": False,
        "context_management": "enabled",
        "duplicate_handling": "context_specific",
        "extraction_available": True
    }

@router.get("/upload/sessions")
async def get_upload_sessions():
    """Debug endpoint to see stored sessions with context information"""
    try:
        session_info = {}
        for session_id, docs in simple_document_store.items():
            # Group documents by context
            contexts = {}
            for doc in docs:
                context_id = doc.get('context_id', 'default')
                if context_id not in contexts:
                    contexts[context_id] = []
                contexts[context_id].append({
                    "filename": doc["filename"],
                    "chunks": doc["chunk_count"],
                    "size": doc["file_size"],
                    "type": doc["file_type"],
                    "upload_time": doc.get("upload_timestamp", "unknown")
                })
            
            session_info[session_id] = {
                "total_documents": len(docs),
                "contexts": contexts,
                "context_count": len(contexts)
            }
        
        return {
            "total_sessions": len(simple_document_store),
            "sessions": session_info,
            "storage_type": "context_aware"
        }
    except Exception as e:
        return {"error": str(e), "note": "Using context-aware document storage"}

@router.get("/upload/context/{context_id}")
async def get_context_documents(
    context_id: str,
    session_id: Optional[str] = Cookie(default=None)
):
    """Get documents for a specific context"""
    if not session_id or session_id not in simple_document_store:
        raise HTTPException(status_code=404, detail="Session not found")
    
    docs = simple_document_store[session_id]
    context_docs = [
        {
            "filename": doc["filename"],
            "chunks": doc["chunk_count"],
            "size": doc["file_size"],
            "type": doc["file_type"],
            "upload_time": doc.get("upload_timestamp", "unknown"),
            "file_hash": doc["file_hash"][:8] + "..."  # Show partial hash for debugging
        }
        for doc in docs if doc.get('context_id', 'default') == context_id
    ]
    
    return {
        "context_id": context_id,
        "document_count": len(context_docs),
        "documents": context_docs
    }

@router.delete("/upload/context/{context_id}/file/{filename}")
async def delete_context_file(
    context_id: str,
    filename: str,
    session_id: Optional[str] = Cookie(default=None)
):
    """Delete a specific file from a specific context"""
    if not session_id or session_id not in simple_document_store:
        raise HTTPException(status_code=404, detail="Session not found")
    
    docs = simple_document_store[session_id]
    initial_count = len(docs)
    
    # Remove the specific file from the specific context
    simple_document_store[session_id] = [
        doc for doc in docs 
        if not (doc.get('filename') == filename and doc.get('context_id', 'default') == context_id)
    ]
    
    removed_count = initial_count - len(simple_document_store[session_id])
    
    if removed_count == 0:
        raise HTTPException(status_code=404, detail=f"File '{filename}' not found in context '{context_id}'")
    
    return {
        "message": f"Successfully removed '{filename}' from context '{context_id}'",
        "removed_count": removed_count,
        "remaining_documents": len(simple_document_store[session_id])
    }

@router.post("/upload/cleanup")
async def cleanup_temp_files():
    """Clean up any leftover temp files"""
    import os
    import glob
    
    temp_files = glob.glob("temp_upload_*")
    cleaned = []
    
    for temp_file in temp_files:
        try:
            os.remove(temp_file)
            cleaned.append(temp_file)
        except Exception as e:
            logger.warning(f"Could not remove {temp_file}: {e}")
    
    return {
        "message": f"Cleaned up {len(cleaned)} temp files",
        "cleaned_files": cleaned,
        "note": "Context-aware upload system doesn't create temp files"
    }