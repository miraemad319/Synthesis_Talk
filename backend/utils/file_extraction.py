import os
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any
from pdfminer.high_level import extract_text
from docx import Document
import json


class FileExtractionError(Exception):
    """Custom exception for file extraction errors"""
    pass


def get_file_info(file_path: str) -> Dict[str, Any]:
    """Get metadata about a file"""
    try:
        path = Path(file_path)
        stat = path.stat()
        mime_type, _ = mimetypes.guess_type(file_path)
        
        return {
            "filename": path.name,
            "size": stat.st_size,
            "extension": path.suffix.lower(),
            "mime_type": mime_type,
            "created": stat.st_ctime,
            "modified": stat.st_mtime
        }
    except Exception as e:
        raise FileExtractionError(f"Error getting file info: {e}")


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file using pdfminer"""
    try:
        if not os.path.exists(file_path):
            raise FileExtractionError(f"File not found: {file_path}")
        
        text = extract_text(file_path)
        if not text or text.strip() == "":
            raise FileExtractionError("PDF appears to be empty or contains no extractable text")
        
        return text.strip()
    except FileExtractionError:
        raise
    except Exception as e:
        raise FileExtractionError(f"Error extracting PDF text: {e}")


def extract_text_from_txt(file_path: str) -> str:
    """Read plain text from a .txt file with encoding detection"""
    try:
        if not os.path.exists(file_path):
            raise FileExtractionError(f"File not found: {file_path}")
        
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                    if content:
                        return content.strip()
            except UnicodeDecodeError:
                continue
        
        raise FileExtractionError("Could not decode text file with any supported encoding")
        
    except FileExtractionError:
        raise
    except Exception as e:
        raise FileExtractionError(f"Error reading TXT file: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word (.docx) file using python-docx"""
    try:
        if not os.path.exists(file_path):
            raise FileExtractionError(f"File not found: {file_path}")
        
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraphs + table_text
        
        if not all_text:
            raise FileExtractionError("DOCX file appears to be empty")
        
        return "\n".join(all_text)
        
    except FileExtractionError:
        raise
    except Exception as e:
        raise FileExtractionError(f"Error reading DOCX file: {e}")


def extract_text_from_json(file_path: str) -> str:
    """Extract text content from JSON files"""
    try:
        if not os.path.exists(file_path):
            raise FileExtractionError(f"File not found: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Convert JSON to readable text format
        if isinstance(data, dict):
            text_parts = []
            for key, value in data.items():
                if isinstance(value, (str, int, float)):
                    text_parts.append(f"{key}: {value}")
                elif isinstance(value, list):
                    text_parts.append(f"{key}: {', '.join(map(str, value))}")
                else:
                    text_parts.append(f"{key}: {str(value)}")
            return "\n".join(text_parts)
        else:
            return str(data)
            
    except FileExtractionError:
        raise
    except Exception as e:
        raise FileExtractionError(f"Error reading JSON file: {e}")


def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Extract text from various file formats
    
    Returns:
        Dict containing:
        - text: extracted text content
        - metadata: file information
        - success: boolean indicating success
        - error: error message if failed
    """
    result = {
        "text": "",
        "metadata": {},
        "success": False,
        "error": None
    }
    
    try:
        # Get file metadata
        result["metadata"] = get_file_info(file_path)
        extension = result["metadata"]["extension"]
        
        # Extract text based on file type
        if extension == ".pdf":
            result["text"] = extract_text_from_pdf(file_path)
        elif extension == ".txt":
            result["text"] = extract_text_from_txt(file_path)
        elif extension in [".docx", ".doc"]:
            if extension == ".doc":
                raise FileExtractionError("Legacy .doc files not supported. Please convert to .docx format.")
            result["text"] = extract_text_from_docx(file_path)
        elif extension == ".json":
            result["text"] = extract_text_from_json(file_path)
        else:
            raise FileExtractionError(f"Unsupported file format: {extension}")
        
        result["success"] = True
        
    except FileExtractionError as e:
        result["error"] = str(e)
    except Exception as e:
        result["error"] = f"Unexpected error: {e}"
    
    return result


def get_supported_extensions() -> list[str]:
    """Get list of supported file extensions"""
    return [".pdf", ".txt", ".docx", ".json"]


def is_supported_file(file_path: str) -> bool:
    """Check if file type is supported"""
    extension = Path(file_path).suffix.lower()
    return extension in get_supported_extensions()


# Legacy functions for backward compatibility
def extract_text_from_pdf_legacy(file_path):
    """Legacy function - kept for backward compatibility"""
    try:
        return extract_text_from_pdf(file_path)
    except FileExtractionError as e:
        return f"Error extracting PDF text: {e}"


def extract_text_from_txt_legacy(file_path):
    """Legacy function - kept for backward compatibility"""
    try:
        return extract_text_from_txt(file_path)
    except FileExtractionError as e:
        return f"Error reading TXT file: {e}"


def extract_text_from_docx_legacy(file_path):
    """Legacy function - kept for backward compatibility"""
    try:
        return extract_text_from_docx(file_path)
    except FileExtractionError as e:
        return f"Error reading DOCX file: {e}"
